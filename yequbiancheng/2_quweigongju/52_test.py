# 
import requests
import json
from pptx import Presentation
import docx
import time
import collections.abc

path = "C:/Users/ChemYS/learngit/yequbiancheng/2_quweigongju/52/ted.docx"
docxFile = docx.Document(path)
paragraphs = docxFile.paragraphs
docxNew = docx.Document()
#print(paragraph)
for pra in paragraphs:
    pra_text = pra.text
    if pra_text != "":
        time.sleep(1)
        word = {"i":pra_text,"doctype":"json"}
        url = "https://fanyi.youdao.com/translate"
        header = {"User-Agent": "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_14_6) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/84.0.4147.89 Safari/537.36"}
        response = requests.post(url,data=word,headers=header)
        res_text = response.text
        data = json.loads(res_text)
        #print(data)
        """translate1 =  data["translateResult"][0][0]["src"]
        translate2 = data["translateResult"][0][0]["tgt"]
        pra.text = translate1 +"\n" + translate2"""
        docxNew.add_paragraph(data["translateResult"][0][0]["src"])
        docxNew.add_paragraph(data["translateResult"][0][0]["tgt"])
 # TODO 将pptxFile保存至/yequ/new_lesson.pptx                     
docxNew.save("C:/Users/ChemYS/learngit/yequbiancheng/2_quweigongju/52/new_learn.docx")